from flask import Flask, request, render_template_string, send_file
import os
from openai import OpenAI
import PyPDF2
import pandas as pd
import base64
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import io

app = Flask(__name__)

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

historial = []

roles = {
"formulador":"Eres experto en formulación de proyectos.",
"presupuestos":"Eres ingeniero experto en presupuestos de obra.",
"licitaciones":"Eres experto en contratación estatal.",
"legal":"Eres abogado administrativo colombiano.",
"financiero":"Eres experto en evaluación financiera.",
"diagnostico":"Eres consultor organizacional."
}

# ================= DASHBOARD =================

@app.route("/")
def panel():
    return render_template_string("""

<html>
<head>
<title>FUNCREDES IA</title>
<style>

body{margin:0;font-family:Arial;background:#f1f4f9;}

.sidebar{position:fixed;width:240px;height:100%;background:#0b2545;color:white;padding:20px;}

.logo{font-size:22px;font-weight:bold;margin-bottom:30px;}

.content{margin-left:260px;padding:40px;}

.cards{display:grid;grid-template-columns:repeat(auto-fit,minmax(230px,1fr));gap:25px;}

.card{background:white;padding:35px;border-radius:15px;box-shadow:0 6px 18px rgba(0,0,0,0.1);
font-size:18px;font-weight:bold;cursor:pointer;text-align:center;transition:0.3s;}

.card:hover{transform:translateY(-5px);}

a{text-decoration:none;color:black;}

</style>
</head>

<body>

<div class="sidebar">
<div class="logo">FUNCREDES IA</div>
</div>

<div class="content">

<h1>Panel de Agentes Inteligentes</h1>

<div class="cards">

<a href="/chat/formulador"><div class="card">🤖 Formular Proyecto</div></a>
<a href="/chat/presupuestos"><div class="card">💰 Presupuesto Obra</div></a>
<a href="/chat/licitaciones"><div class="card">📑 Analizar Licitación</div></a>
<a href="/chat/legal"><div class="card">⚖️ Consulta Legal</div></a>
<a href="/chat/financiero"><div class="card">📊 Evaluación Financiera</div></a>
<a href="/chat/diagnostico"><div class="card">📋 Diagnóstico</div></a>

</div>

</div>

</body>
</html>

""")

# ================= CHAT =================

@app.route("/chat/<agente>", methods=["GET","POST"])
def chat(agente):

    respuesta=""
    imagen_base64=None
    texto_docs=""

    if request.method=="POST":

        consulta=request.form["consulta"]
        archivos = request.files.getlist("files")

        for archivo in archivos:

            if archivo.filename.endswith(".pdf"):
                lector = PyPDF2.PdfReader(archivo)
                for p in lector.pages:
                    texto_docs += p.extract_text() or ""

            elif archivo.filename.endswith(".xlsx"):
                df = pd.read_excel(archivo)
                texto_docs += df.to_string()

            elif archivo.filename.lower().endswith((".jpg",".jpeg",".png")):
                imagen_bytes = archivo.read()
                imagen_base64 = base64.b64encode(imagen_bytes).decode()

        if imagen_base64:

            completion = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role":"system","content":roles[agente]},
                    {"role":"user","content":[
                        {"type":"text","text":consulta + texto_docs},
                        {"type":"image_url","image_url":{"url":f"data:image/jpeg;base64,{imagen_base64}"}}
                    ]}
                ]
            )

        else:

            prompt = roles[agente] + f"""
Consulta:
{consulta}

Contenido documentos:
{texto_docs}
"""

            completion = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[{"role":"user","content":prompt}]
            )

        respuesta = completion.choices[0].message.content
        historial.append((consulta,respuesta))

    return render_template_string("""

<html>
<head>

<style>

body{margin:0;font-family:Arial;background:#0f172a;color:white;}

.top{background:#020617;padding:15px;font-size:22px;font-weight:bold;}

.layout{display:flex;height:88vh;}

.chat{flex:3;padding:20px;overflow:auto;}

.historial{flex:1;background:#020617;padding:15px;overflow:auto;}

.msg-user{background:#2563eb;padding:12px;border-radius:10px;margin:10px;text-align:right;}

.msg-ia{background:#1e293b;padding:12px;border-radius:10px;margin:10px;}

.input{position:fixed;bottom:0;width:100%;background:#020617;padding:15px;}

textarea{width:60%;height:60px;border-radius:8px;}

button{padding:10px 25px;border-radius:8px;background:#2563eb;color:white;border:none;}

.preview img{max-width:120px;margin:5px;}

</style>

<script>

function previewImages(){
let preview=document.getElementById('preview')
preview.innerHTML=""
let files=document.getElementById('files').files

for(let i=0;i<files.length;i++){
let reader=new FileReader()
reader.onload=function(e){
preview.innerHTML+="<img src='"+e.target.result+"'>"
}
reader.readAsDataURL(files[i])
}
}

</script>

</head>

<body>

<div class="top">
Agente {{agente}}
<a href="/" style="color:white;margin-left:20px">⬅ Panel</a>
</div>

<div class="layout">

<div class="chat">

{% for c,r in historial %}
<div class="msg-user">{{c}}</div>
<div class="msg-ia">{{r}}</div>
{% endfor %}

</div>

<div class="historial">
<h3>Historial</h3>
{% for c,r in historial %}
<p>Consulta enviada</p>
{% endfor %}
</div>

</div>

<div class="input">

<form method="post" enctype="multipart/form-data">

<textarea name="consulta"></textarea>

<input type="file" id="files" name="files" multiple onchange="previewImages()">

<button>Enviar</button>

<a href="/word"><button type="button">📄 Word</button></a>
<a href="/pdf"><button type="button">📑 PDF</button></a>

<div id="preview" class="preview"></div>

</form>

</div>

</body>
</html>

""", historial=historial, agente=agente.upper())

# ================= WORD =================

@app.route("/word")
def word():
    texto = historial[-1][1]
    doc = Document()
    doc.add_heading("Informe IA",0)
    doc.add_paragraph(texto)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer,as_attachment=True,download_name="informe.docx")

# ================= PDF =================

@app.route("/pdf")
def pdf():
    texto = historial[-1][1]
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = [Paragraph(texto, styles["Normal"])]
    doc.build(story)
    buffer.seek(0)
    return send_file(buffer,as_attachment=True,download_name="informe.pdf")

if __name__ == "__main__":
    app.run()
